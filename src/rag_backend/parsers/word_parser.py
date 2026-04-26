# -*- coding: utf-8 -*-

from pathlib import Path
from .base import DocumentParser

class WordParser(DocumentParser):
    def is_supported(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".docx"

    def parse(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        content = []
        current_section = ""
        for para in doc.paragraphs:
            if para.text.strip():
                is_title = False
                heading_level = 0
                if para.style.name.startswith('Heading'):
                    try:
                        heading_level = int(para.style.name.split(' ')[1])
                        is_title = True
                    except (IndexError, ValueError):
                        pass
                if is_title:
                    current_section = para.text
                    content.append(f"## {para.text}")
                else:
                    content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                content.append(" | ".join(row_data))
        return "\n".join(content)