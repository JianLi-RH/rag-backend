# -*- coding: gbk -*-

# parser factory class
# create a parser instance based on .env config
# parser will get plain text content from the file
# parser will return the plain text content as a string
# parser will handle different file types

from abc import ABC, abstractmethod
from pathlib import Path
import pdfplumber

from logger import logger

class DocumentParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> str:
        pass

    @abstractmethod
    def is_supported(self, file_path: str) -> bool:
        pass

class WordParser(DocumentParser):
    def is_supported(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".docx"

    def parse(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        content = []
        current_section = ""
        for para in doc.paragraphs:
            if para.text.strip():
                is_title = False
                heading_level = 0
                if para.style.name.startswith('Heading'):
                    try:
                        heading_level = int(para.style.name.split(' ')[1])
                        is_title = True
                    except (IndexError, ValueError):
                        pass
                if is_title:
                    current_section = para.text
                    content.append(f"## {para.text}")
                else:
                    content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                content.append(" | ".join(row_data))
        return "\n".join(content)


class PDFParser(DocumentParser):
    def is_supported(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".pdf"

    def parse(self, file_path: str) -> str:
        with pdfplumber.open(file_path) as pdf:
            # ������ȡҳ���ı�
            page_texts = [page.extract_text() or "" for page in pdf.pages]
            all_text = " ".join(page_texts)
            return all_text


class MarkdownParser(DocumentParser):
    def is_supported(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".md"

    def parse(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            md_content = f.read()
        return md_content


class TextParser(DocumentParser):
    def is_supported(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".txt"

    def parse(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return text


class ExcelParser(DocumentParser):
    def is_supported(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".xlsx"

    def parse(self, file_path: str) -> str:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True)
        content = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            content.append(f"=== Sheet: {sheet_name} ===")
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                content.append(" | ".join(row_data))
        return "\n".join(content)


class DocumentParserFactory:
    _parsers = {
        ".docx": WordParser,
        ".pdf": PDFParser,
        ".md": MarkdownParser,
        ".txt": TextParser,
        ".xlsx": ExcelParser,
    }

    @classmethod
    def get_parser(cls, file_path: str) -> DocumentParser:
        ext = Path(file_path).suffix.lower()
        parser_class = cls._parsers.get(ext)
        if parser_class is None:
            raise ValueError(f"Unsupported file type: {ext}")
        return parser_class()

    @classmethod
    def get_supported_extensions(cls) -> list:
        return list(cls._parsers.keys())

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in cls._parsers


if __name__ == "__main__":
    abs_path = r"D:\code\AI\rag_backend\uploaded_docs\python-3.13-docs-pdf-a4\docs-pdf\library.pdf"
    import datetime
    print(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    parser = DocumentParserFactory.get_parser(abs_path)
    parsed_content = parser.parse(abs_path)
    print(parsed_content[:10])
    print(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
